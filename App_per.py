import io
import re
from typing import Dict, List

import streamlit as st
import pdfplumber
from docx import Document
import pandas as pd
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity


# ═══════════════════════════════════════════════════════
# PDF EXTRACTION — POSITIONAL WORD GROUPING
# ═══════════════════════════════════════════════════════

def get_words_with_position(page) -> List[dict]:
    """Return all words with x0, y0, x1, y1 from pdfplumber page."""
    return page.extract_words(
        x_tolerance=3,
        y_tolerance=3,
        keep_blank_chars=False,
        use_text_flow=False,
    )


def group_words_into_lines(words: List[dict], y_tolerance: int = 5) -> List[str]:
    """
    Group words that share approximately the same y0 (same row)
    and sort them left-to-right. Returns list of line strings.
    """
    if not words:
        return []
    
    # Sort by top (y0) then left (x0)
    words_sorted = sorted(words, key=lambda w: (round(float(w["top"]) / y_tolerance), float(w["x0"])))
    
    lines = []
    current_line_words = [words_sorted[0]]
    current_top = round(float(words_sorted[0]["top"]) / y_tolerance)
    
    for word in words_sorted[1:]:
        word_top = round(float(word["top"]) / y_tolerance)
        if word_top == current_top:
            current_line_words.append(word)
        else:
            lines.append(" ".join(w["text"] for w in current_line_words))
            current_line_words = [word]
            current_top = word_top
    
    lines.append(" ".join(w["text"] for w in current_line_words))
    return lines


def extract_pdf_as_lines(file_bytes: bytes) -> List[str]:
    """
    Extract full PDF as positionally-sorted lines.
    Each line contains words that appear on same y-coordinate,
    ordered left to right — so table row content is on ONE line.
    """
    all_lines = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            words = get_words_with_position(page)
            lines = group_words_into_lines(words, y_tolerance=5)
            all_lines.extend(lines)
    return all_lines


def parse_pdf_lines(lines: List[str]) -> Dict[str, str]:
    """
    Parse key fields from positionally-reconstructed lines.
    Since words on the same row are now on the same line, patterns like:
      'Title Investment Banking Supervisory Structure Gaps Issue ID ISSUE-00077693'
    are reliably parseable.
    """
    fields = {
        "Title": "",
        "Issue ID": "",
        "Description": "",
        "Issue Impact": "",
        "Issue Root Cause": "",
    }

    full_text = "\n".join(lines)

    # ---- Title ----
    # Line looks like: "Title Investment Banking Supervisory Structure Gaps Issue ID ISSUE-00077693"
    # OR: "Title Investment Banking Supervisory Structure Gaps"
    title_match = re.search(
        r"\bTitle\b\s+(.+?)\s+Issue\s+ID\b",
        full_text,
        flags=re.DOTALL,
    )
    if title_match:
        fields["Title"] = " ".join(title_match.group(1).split())
    else:
        # fallback: Title spans to end of line
        for line in lines:
            if re.match(r"^Title\s+\S", line):
                fields["Title"] = re.sub(r"^Title\s+", "", line).strip()
                break

    # ---- Issue ID ----
    issue_id_match = re.search(
        r"\bIssue\s+ID\s+(ISSUE-\S+)",
        full_text,
    )
    if issue_id_match:
        fields["Issue ID"] = issue_id_match.group(1).strip()

    # ---- Description ----
    desc_match = re.search(
        r"Description\s+(.+?)Issue Impact",
        full_text,
        flags=re.DOTALL,
    )
    if desc_match:
        fields["Description"] = " ".join(desc_match.group(1).split())

    # ---- Issue Impact ----
    impact_match = re.search(
        r"Issue Impact\s+(.+?)Issue Root Cause",
        full_text,
        flags=re.DOTALL,
    )
    if impact_match:
        fields["Issue Impact"] = " ".join(impact_match.group(1).split())

    # ---- Issue Root Cause ----
    root_match = re.search(
        r"Issue Root Cause\s+(.+?)Overall Issue Rating",
        full_text,
        flags=re.DOTALL,
    )
    if root_match:
        fields["Issue Root Cause"] = " ".join(root_match.group(1).split())

    return fields


# ═══════════════════════════════════════════════════════
# DOCX EXTRACTION — TABLE CELL WALK (WORKING)
# ═══════════════════════════════════════════════════════

def parse_icp_docx_from_file(file_bytes: bytes) -> Dict[str, str]:
    fields = {
        "Title": "",
        "Issue ID": "",
        "Description": "",
        "Issue Impact": "",
        "Issue Root Cause": "",
    }

    doc = Document(io.BytesIO(file_bytes))
    kv_map = {}

    # Walk table cells row by row
    for table in doc.tables:
        for row in table.rows:
            cells = [" ".join(cell.text.split()) for cell in row.cells]
            i = 0
            while i < len(cells) - 1:
                key = cells[i].rstrip(":").strip()
                value = cells[i + 1].strip()
                if key:
                    kv_map[key] = value
                i += 2

    # Normalize keys
    normalized_kv = {" ".join(k.split()).lower(): v for k, v in kv_map.items()}

    fields["Title"] = normalized_kv.get("issue title", "")
    fields["Issue ID"] = normalized_kv.get("source system issue reference", "")

    # Paragraph-style fields from paragraphs + table cells
    para_text = "\n".join(p.text for p in doc.paragraphs)
    table_text = "\n".join(
        cell.text for table in doc.tables for row in table.rows for cell in row.cells
    )
    combined = para_text + "\n" + table_text

    desc_match = re.search(
        r"Issue Description:\s*(.+?)Issue Root Cause:", combined, flags=re.DOTALL
    )
    if desc_match:
        fields["Description"] = " ".join(desc_match.group(1).split())

    root_match = re.search(
        r"Issue Root Cause:\s*(.+?)Issue Impact:", combined, flags=re.DOTALL
    )
    if root_match:
        fields["Issue Root Cause"] = " ".join(root_match.group(1).split())

    impact_match = re.search(
        r"Issue Impact:\s*(.+?)Background Context:", combined, flags=re.DOTALL
    )
    if impact_match:
        fields["Issue Impact"] = " ".join(impact_match.group(1).split())
    else:
        impact_match2 = re.search(
            r"Issue Impact:\s*(.+?)Section C:", combined, flags=re.DOTALL
        )
        if impact_match2:
            fields["Issue Impact"] = " ".join(impact_match2.group(1).split())

    return fields


# ═══════════════════════════════════════════════════════
# SIMILARITY COMPUTATION
# ═══════════════════════════════════════════════════════

def text_similarity(a: str, b: str) -> float:
    if not a or not b:
        return 0.0
    vectorizer = TfidfVectorizer().fit([a, b])
    tfidf = vectorizer.transform([a, b])
    return float(cosine_similarity(tfidf[0:1], tfidf[1:2])[0][0])


def compute_similarity(f1: Dict[str, str], f2: Dict[str, str]) -> Dict[str, float]:
    scores = {}
    scores["Issue ID"] = 1.0 if f1.get("Issue ID") == f2.get("Issue ID") else 0.0
    scores["Title"] = text_similarity(f1.get("Title", ""), f2.get("Title", ""))
    scores["Description"] = text_similarity(f1.get("Description", ""), f2.get("Description", ""))
    scores["Issue Root Cause"] = text_similarity(f1.get("Issue Root Cause", ""), f2.get("Issue Root Cause", ""))
    scores["Issue Impact"] = text_similarity(f1.get("Issue Impact", ""), f2.get("Issue Impact", ""))
    scores["Overall"] = sum(scores.values()) / len(scores)
    return scores


def match_label(score: float, threshold: float) -> str:
    return "✅ Match" if score >= threshold else "❌ Mismatch"


# ═══════════════════════════════════════════════════════
# STREAMLIT UI
# ═══════════════════════════════════════════════════════

def main():
    st.set_page_config(page_title="IBF vs ICP Comparison", layout="wide")
    st.title("📄 IBF vs ICP Comparison")
    st.write("Upload the Issue Briefing Form (PDF) and Issue Closure Pack (DOCX) to compare key fields.")

    col1, col2 = st.columns(2)
    with col1:
        pdf_file = st.file_uploader("Upload Issue Briefing Form (PDF)", type=["pdf"])
    with col2:
        docx_file = st.file_uploader("Upload Issue Closure Pack ICP (DOCX)", type=["docx"])

    threshold = st.slider("Match threshold", 0.0, 1.0, 0.8, 0.05)

    if pdf_file and docx_file and st.button("🔍 Compare"):

        pdf_bytes = pdf_file.read()
        docx_bytes = docx_file.read()

        # --- Parse PDF using positional word grouping ---
        pdf_lines = extract_pdf_as_lines(pdf_bytes)
        f1 = parse_pdf_lines(pdf_lines)

        # --- Parse DOCX using table cell walk ---
        f2 = parse_icp_docx_from_file(docx_bytes)

        # --- DEBUG expanders ---
        with st.expander("🔍 PDF reconstructed lines (debug)"):
            for i, line in enumerate(pdf_lines[:40]):
                st.text(f"[{i}] {line}")

        with st.expander("🔍 Parsed fields from IBF (PDF)"):
            st.json(f1)

        with st.expander("🔍 Parsed fields from ICP (DOCX)"):
            st.json(f2)

        # --- Key-value comparison table ---
        st.subheader("📋 Key Fields Comparison")
        rows_kv = [
            {"Attribute": "Title",            "IBF (PDF)": f1.get("Title", ""),            "ICP (DOCX)": f2.get("Title", "")},
            {"Attribute": "Issue ID",         "IBF (PDF)": f1.get("Issue ID", ""),         "ICP (DOCX)": f2.get("Issue ID", "")},
            {"Attribute": "Description",      "IBF (PDF)": f1.get("Description", ""),      "ICP (DOCX)": f2.get("Description", "")},
            {"Attribute": "Issue Root Cause", "IBF (PDF)": f1.get("Issue Root Cause", ""), "ICP (DOCX)": f2.get("Issue Root Cause", "")},
            {"Attribute": "Issue Impact",     "IBF (PDF)": f1.get("Issue Impact", ""),     "ICP (DOCX)": f2.get("Issue Impact", "")},
        ]
        df_kv = pd.DataFrame(rows_kv)
        st.dataframe(df_kv, use_container_width=True)

        csv_kv = df_kv.to_csv(index=False).encode("utf-8")
        st.download_button("⬇️ Download key fields CSV", csv_kv, "ibf_icp_key_fields.csv", "text/csv")

        # --- Similarity scores table ---
        st.subheader("📊 Similarity Scores")
        scores = compute_similarity(f1, f2)
        sim_rows = []
        for field in ["Issue ID", "Title", "Description", "Issue Root Cause", "Issue Impact"]:
            score = scores[field]
            thresh = 1.0 if field == "Issue ID" else threshold
            sim_rows.append({
                "Field": field,
                "Similarity Score": round(score, 3),
                "Result": match_label(score, thresh),
            })
        sim_rows.append({
            "Field": "Overall",
            "Similarity Score": round(scores["Overall"], 3),
            "Result": match_label(scores["Overall"], threshold),
        })
        df_sim = pd.DataFrame(sim_rows)
        st.dataframe(df_sim, use_container_width=True)

        csv_sim = df_sim.to_csv(index=False).encode("utf-8")
        st.download_button("⬇️ Download similarity scores CSV", csv_sim, "ibf_icp_similarity_scores.csv", "text/csv")


if __name__ == "__main__":
    main()
